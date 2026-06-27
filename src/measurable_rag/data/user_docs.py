"""Load user-uploaded documents — the 'bring your own corpus' (NotebookLM-style) path.

Turns uploaded files into Document objects so the SAME pipeline (chunking with
offsets -> embedding -> retrieval -> citation verification -> abstention) runs
over them.

Note the distinction from the SciFact benchmark: SciFact has relevance labels,
so it's what we use to MEASURE retrieval (Recall@k etc.). User documents have no
labels — this path is the interactive demo, not the measurement.
"""
from __future__ import annotations

import io

from .models import Document


def extract_text(name: str, data: bytes) -> str:
    """Extract plain text from an uploaded file.

    Supports PDF, Word (.docx), Excel (.xlsx), and any UTF-8 text format
    (.txt/.md/.csv/.tsv/.json and source code). Unknown types fall back to a
    UTF-8 decode, so most text-based files just work.
    """
    lower = name.lower()
    if lower.endswith(".pdf"):
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        return "\n\n".join((page.extract_text() or "") for page in reader.pages).strip()
    if lower.endswith(".docx"):
        from docx import Document as DocxDocument

        doc = DocxDocument(io.BytesIO(data))
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip()).strip()
    if lower.endswith((".xlsx", ".xlsm")):
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        lines: list[str] = []
        for ws in wb.worksheets:
            lines.append(f"# Sheet: {ws.title}")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    lines.append(" | ".join(cells))
        return "\n".join(lines).strip()
    # csv / tsv / txt / md / json / source code / anything else text-like
    return data.decode("utf-8", errors="replace").strip()


def load_documents(payloads: list[tuple[str, bytes]]) -> dict[str, Document]:
    """Build {doc_id: Document} from (filename, bytes) pairs.

    One Document per file; its filename is the doc id and offsets index into the
    file's extracted text, so citation spans point back into the right file.
    """
    docs: dict[str, Document] = {}
    for name, data in payloads:
        text = extract_text(name, data)
        if text:
            docs[name] = Document(doc_id=name, text=text, title=name)
    return docs
